import os
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path.cwd()
SOURCE = ROOT / "docs" / "csb-migration-guide.md"
OUT_DIR = ROOT / "outputs" / "csb-migration"
OUT_PATH = Path(os.environ.get("CSB_GUIDE_DOCX_OUT", OUT_DIR / "arabian-drilling-csb-migration-guide.docx"))

GREEN = "003C34"
ORANGE = "FF5A00"
GOLD = "E5A93F"
LIGHT = "F4F6F5"
PALE_GOLD = "FBF3DF"
BORDER = "D8E0DD"
TEXT = "24312F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=110, start=110, bottom=110, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_paragraph(paragraph, font_size=10.5, color=TEXT, bold=False):
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def add_paragraph(doc, text="", style=None):
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text)
    style_paragraph(p)
    return p


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(GREEN)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_callout(doc, title, body, fill=PALE_GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=GOLD, size="8")
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(GREEN)
    p.runs[0].font.size = Pt(11)
    body_p = cell.add_paragraph()
    add_inline_runs(body_p, body)
    style_paragraph(body_p, font_size=10)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        set_cell_shading(cell, GREEN)
        set_cell_border(cell, color=GREEN)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(paragraph, font_size=9.5, color="FFFFFF", bold=True)
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.text = value
            set_cell_shading(cell, "FFFFFF" if row_index % 2 == 0 else LIGHT)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                style_paragraph(paragraph, font_size=9.2)
    doc.add_paragraph()


def build_doc():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)
    for name, size, color in [
        ("Heading 1", 18, GREEN),
        ("Heading 2", 14, GREEN),
        ("Heading 3", 11.5, ORANGE),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Arabian Drilling Careers")
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("SuccessFactors Career Site Builder Migration Guide")
    srun.font.name = "Aptos"
    srun.font.size = Pt(15)
    srun.font.color.rgb = RGBColor.from_string(ORANGE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mrun = meta.add_run("Prepared as a step-by-step playbook for moving the prototype into CSB | 2026-05-15")
    mrun.font.size = Pt(9.5)
    mrun.font.color.rgb = RGBColor.from_string(TEXT)

    add_callout(
        doc,
        "Migration Principle",
        "Build with native Career Site Builder components first. Use Custom Plugin only for sections CSB cannot express cleanly, and keep those plugins isolated from native search, job detail, candidate profile, and apply flows.",
    )

    lines = markdown.splitlines()
    i = 1
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_callout(doc, "Code / Pattern", "\n".join(code_lines), fill=LIGHT)
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
            style_paragraph(p)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\. ", "", line))
            style_paragraph(p)
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for table_line in table_lines:
                cells = [c.strip() for c in table_line.strip("|").split("|")]
                if all(re.fullmatch(r"-+", c.replace(" ", "")) for c in cells):
                    continue
                rows.append(cells)
            if rows:
                add_table(doc, rows[0], rows[1:])
            continue
        else:
            add_paragraph(doc, line)
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Arabian Drilling Careers | CSB Migration Guide")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(GREEN)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
